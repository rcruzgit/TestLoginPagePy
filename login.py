import os
import sys
from datetime import datetime
import base64
from encodings.base64_codec import base64_decode
from time import sleep
from behave import *
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions
from selenium.webdriver.support.wait import WebDriverWait
from selenium.webdriver.common.keys import Keys
import pyautogui
from docx import Document
from docx.shared import Inches
from PIL import ImageGrab

# Get current directory path
current_directory = os.getcwd()
# Get timestamp and format to use later
now = datetime.now()
timestamp = now.strftime("%Y%m%d_%H%M%S")
# format filename with timestamp
fname = f"TestResult_{timestamp}.docx"
# set filepath
newfile = os.path.join(current_directory, fname)
tmpshot = os.path.join(current_directory, "tmp.png")
# create .docx file
document = Document()
document.save(newfile)
print(f"Files '{newfile}' and '{tmpshot}' created successfully at '{current_directory}'.")

def screencap(txt):
    # capture screen
    shot = pyautogui.screenshot()  # Take a screenshot full screen
    shot.save(tmpshot)  # Save the screenshot
    # Append to the document. Doc must exist.
    doc = Document(newfile)  # Open the document
    doc.add_paragraph(txt)
    doc.add_picture(tmpshot, width=Inches(6))  # Add the image, 7 inches wide
    doc.save(newfile)  # Update the document

def highlight(element,txt):
    driver = element._parent
    def apply_style(s):
        driver.execute_script("arguments[0].setAttribute('style', arguments[1]);", element, s)
    original_style = element.get_attribute('style')
    i = 0
    while i < 2:
        apply_style("background: yellow; border: 2px solid red;")
        sleep(0.2)
        if(i==1):
            screencap(txt) #capture screenshot while element is highlighted
        apply_style(original_style)
        sleep(0.2)
        i = i + 1

@given('user launched browser')
def launchbrowser(context):
    # setup chrome for opening the page
    options = webdriver.ChromeOptions()
    # setup options for the current session
    options.add_experimental_option('excludeSwitches', ['enable-logging'])
    # launch chrome and open target URL
    context.driver = webdriver.Chrome(options=options)
    sleep(2)
    txtval = 'Given user launched browser'
    screencap(txtval)

@given(u'user opened the page')
def openpage(context):
    URL = "https://practicetestautomation.com/practice-test-login/"
    context.driver.get(URL)
    context.driver.maximize_window()
    sleep(2)
    txtval = 'And user opened the page'
    screencap(txtval)


@given(u'user entered username "{username}"')
def enterusername(context, username):
    # Scroll to the username field
    wait = WebDriverWait(context.driver, 10)
    element = wait.until(expected_conditions.visibility_of_element_located((By.XPATH, "//input[@id='username']")))
    context.driver.execute_script("arguments[0].scrollIntoView(true);", element)
    # Enter values in the username field
    element.send_keys(username)
    txtval = f'And user entered username "{username}"'
    highlight(element,txtval)
    sleep(2)

@given(u'user entered password "{pwd}"')
def enterpwd(context, pwd):
    #Navigate to Password field
    element = context.driver.find_element(By.XPATH, "//input[@id='password']")
    # Decode password
    decoded_pwd = base64.b64decode(pwd).decode('utf-8')
    # Enter password value
    element.send_keys(decoded_pwd)
    txtval = f'And user entered password "{pwd}"'
    highlight(element,txtval)
    sleep(2)

@when(u'clicked on Submit button')
def clicksubmit(context):
    # Locate and click on the submit button
    element = context.driver.find_element(By.XPATH, "//button[@id='submit']")
    txtval = 'When user clicked Submit button'
    highlight(element,txtval)
    element.click()
    sleep(2)

@then(u'new page URL will be displayed')
def checkurl(context):
    #Compare the expected URL redirect
    targetURL = 'https://practicetestautomation.com/logged-in-successfully/'
    assert targetURL==context.driver.current_url
    sleep(2)
    txtval = 'Then new page URL will be displayed'
    screencap(txtval)

@then(u'new page contains expected text "{text}"')
def checkpagetext(context, text):
    # Locate the text identifier in the page for successful login
    wait = WebDriverWait(context.driver, 10)
    element = wait.until(expected_conditions.visibility_of_element_located((By.XPATH, "//h1[text()[contains(.,text)]]")))
    txtval = f"And new page contains expected text '{text}'"
    highlight(element,txtval)

@then(u'Log Out button is displayed')
def checklogout(context):
    # Locate the logout button
    element = context.driver.find_element(By.XPATH, "//a[@class='wp-block-button__link has-text-color has-background has-very-dark-gray-background-color']")
    txtval = 'And Log Out button is displayed'
    highlight(element,txtval)
